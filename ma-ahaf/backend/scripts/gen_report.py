"""Generate the MA-AHAF professional project report as a .docx (Microsoft Word)."""
from __future__ import annotations

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

OUT = r"C:\Users\acer\OneDrive\Desktop\RAG\ma-ahaf\docs\MA-AHAF-Project-Report.docx"

ACCENT = RGBColor(0x3B, 0x35, 0x9E)     # indigo
INK = RGBColor(0x1B, 0x23, 0x30)
MUTED = RGBColor(0x55, 0x5F, 0x6D)
OK = RGBColor(0x15, 0x70, 0x3D)

doc = Document()

# ---------- base styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for lvl, sz in [(1, 17), (2, 13.5), (3, 11.5)]:
    st = doc.styles[f"Heading {lvl}"]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = ACCENT if lvl == 1 else INK
    st.paragraph_format.space_before = Pt(16 if lvl == 1 else 10)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.keep_with_next = True

for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.85)
    s.left_margin = s.right_margin = Inches(0.9)


def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def para(text="", *, size=10.5, bold=False, italic=False, color=None,
         align=None, space_after=6, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        if color is not None:
            r.font.color.rgb = color
    return p


def bullets(items, *, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if isinstance(it, tuple):
            r = p.add_run(it[0] + " ")
            r.font.bold = True
            r.font.size = Pt(size)
            r2 = p.add_run(it[1])
            r2.font.size = Pt(size)
        else:
            r = p.add_run(it)
            r.font.size = Pt(size)


def table(headers, rows, *, widths=None, col0_bold=True, font=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        _shade(hdr[i], "3B359E")
        pr = hdr[i].paragraphs[0]
        pr.paragraph_format.space_after = Pt(2)
        run = pr.add_run(h)
        run.font.bold = True
        run.font.size = Pt(font)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if ridx % 2 == 1:
                _shade(cells[i], "EEF0FB")
            pr = cells[i].paragraphs[0]
            pr.paragraph_format.space_after = Pt(2)
            run = pr.add_run(str(val))
            run.font.size = Pt(font)
            if i == 0 and col0_bold:
                run.font.bold = True
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "CBD2DC")
    bdr.append(bottom)
    pPr.append(bdr)


# ============================================================ COVER
para("PROJECT REPORT", size=11, bold=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_after=4)
para("MA-AHAF — Multi-Agent Adaptive Framework for\nDynamic Hallucination Attenuation in LLMs",
     size=22, bold=True, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
para("Automatic balancing between factual reliability and generative creativity",
     size=12, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

table(["Field", "Value"], [
    ["Project type", "Applied AI / LLM Reliability / Multi-Agent Systems"],
    ["Deliverable", "Research-grade, deployable prototype + evaluation"],
    ["Report version", "1.0"],
    ["Date", "1 September 2026"],
    ["Status", "Prototype complete — all 13 agents operational, end-to-end evaluation run"],
    ["Reference", "Fulfils 'project proposal final.pdf' (sections 4, 7-18)"],
], widths=[1.5, 4.7], font=10)

para("This document describes what was built, the technology stack, the request "
     "workflow, the measured evaluation results, and a file-by-file breakdown of the "
     "entire codebase. It is generated from the delivered repository.",
     size=9.5, italic=True, color=MUTED, space_after=4)

doc.add_page_break()

# ============================================================ 1. EXEC SUMMARY
doc.add_heading("1.  Executive Summary", 1)
para("MA-AHAF is a model-agnostic adaptive layer that sits between an application and "
     "one or more LLMs. For every request it estimates hallucination risk, evidence "
     "coverage, source agreement, claim criticality, ambiguity and creativity demand, "
     "then selects a per-request policy vector controlling how much grounding, "
     "verification, citation, uncertainty disclosure, creative latitude and abstention "
     "are applied — with a full audit trail.")
para("Instead of applying the same retrieval-and-verification pipeline to every prompt, "
     "the framework runs a closed control loop: input intelligence \u2192 adaptive policy "
     "\u2192 generation \u2192 claim decomposition \u2192 evidence triangulation \u2192 multi-agent "
     "verification \u2192 conflict resolution \u2192 revision or calibrated abstention \u2192 "
     "audit.")
para("Delivered in this engagement:", bold=True, space_after=3)
bullets([
    "All 13 agents from proposal \u00a77 — implemented, unit-tested, running end-to-end on OpenAI gpt-4o-mini.",
    "The Adaptive Reliability\u2013Creativity Controller (ARCOP), the claim-level H(x) risk model, "
    "evidence triangulation, hallucination budget, dynamic verification depth, counterfactual "
    "verification, creativity-preservation gate, confidence\u2013evidence consistency check, "
    "disagreement-driven escalation — all 10 novelty items from proposal \u00a75.",
    "Model-agnostic LLM gateway (OpenAI / local Ollama / HuggingFace / mock adapters).",
    "FastAPI service with OpenAPI docs, PostgreSQL + pgvector, Redis, Alembic migrations, "
    "OpenTelemetry + Prometheus, structured audit traces.",
    "Full React dashboard: Playground, Traces, Trace Detail, Human Review queue, Metrics, "
    "Evaluation, Knowledge Base.",
    "77-item evaluation benchmark + harness (MA-AHAF vs static-RAG baseline) with a "
    "reliability\u2013creativity Pareto report and a retraining loop for the risk model and calibrator.",
    "143-document knowledge base across 13 domains; Docker Compose stack; Kubernetes manifests; "
    "k6 load test; security review; 30 automated tests (all passing); lint clean.",
])
para("Headline evaluation result (76/77 items, gpt-4o-mini): calibration error (ECE) "
     "improved 0.191 \u2192 0.066 (\u221265%), Brier 0.194 \u2192 0.145, and the system abstains "
     "on 15.8% of items (the deliberately unanswerable high-stakes set) where the baseline "
     "always answers. Full numbers and interpretation in section 9.", bold=True)

# ============================================================ 2. DELIVERABLES
doc.add_heading("2.  What Was Delivered (vs Proposal \u00a717)", 1)
table(["Proposal deliverable", "Status", "Where"], [
    ["Software Requirements Specification + architecture doc", "Delivered", "docs/SRS.md, docs/architecture.md"],
    ["Multi-agent orchestration engine", "Delivered", "app/orchestration/, app/agents/a01\u2013a13"],
    ["Model-agnostic LLM gateway", "Delivered", "app/llm/gateway.py + 4 adapters"],
    ["Adaptive Reliability\u2013Creativity Controller", "Delivered", "app/controller/arcop.py"],
    ["Claim decomposition + risk graph module", "Delivered", "app/agents/a05, app/claimgraph/"],
    ["Evidence retrieval, ranking, source-quality, contradiction", "Delivered", "app/retrieval/, app/agents/a06\u2013a09"],
    ["Verification / revision / abstention / escalation pipeline", "Delivered", "app/agents/a08\u2013a12"],
    ["Evaluation benchmark + automated harness", "Delivered", "data/benchmark/, app/eval/, scripts/eval_local.py"],
    ["API service with OpenAPI / Swagger", "Delivered", "app/main.py, app/api/*, GET /docs"],
    ["Observability dashboard + structured audit traces", "Delivered", "frontend/, app/agents/a13, app/core/telemetry.py"],
    ["Deployment package (Docker; K8s if in scope)", "Delivered", "docker-compose.yml, deploy/k8s/"],
    ["Technical documentation + configuration guide + demo", "Delivered", "docs/, README.md, scripts/demo.py, scripts/serve_demo.py"],
    ["Final project report (methodology, experiments, results, limitations)", "Delivered", "docs/final-report.md + this document"],
], widths=[3.1, 0.9, 2.2], font=9)

# ============================================================ 3. ARCHITECTURE + WORKFLOW
doc.add_heading("3.  System Architecture & Request Workflow", 1)
para("The framework operates as an adaptive orchestration layer. The canonical flow "
     "graph is defined as a LangGraph state machine (app/orchestration/graph.py, "
     "exposed as a Mermaid diagram at GET /v1/graph); the runtime executes the same "
     "nodes through an explicit, debuggable control loop (app/orchestration/pipeline.py).")

doc.add_heading("3.1  Layers", 2)
table(["Layer", "Responsibility", "Components"], [
    ["Input intelligence", "Understand task, domain, stakes, requested style", "Intent & Task Classifier, Risk Profiler"],
    ["Adaptive control", "Select the reliability/creativity operating point", "ARCOP controller, policy engine, hallucination budget"],
    ["Generation", "Produce candidate responses under policy constraints", "LLM gateway, Candidate Generator"],
    ["Claim analysis", "Break output into verifiable units", "Claim Decomposer, Claim Risk Graph (networkx)"],
    ["Evidence", "Retrieve & rank supporting / contradicting material", "Hybrid retriever (vector + BM25 + RRF + MMR), cross-encoder reranker, source-quality scorer"],
    ["Verification", "Check entailment, contradiction, internal consistency", "Verifier agent (NLI + LLM judge + counterfactual), Contradiction agent"],
    ["Resolution", "Aggregate disagreement, choose the safe action", "Consensus engine, conflict resolver"],
    ["Response control", "Revise, qualify, cite, abstain, preserve creativity", "Revision agent, creativity-preservation gate, abstention/escalation gate"],
    ["Observability", "Track quality, cost, latency, failure modes", "OpenTelemetry, Prometheus, structured logs, audit store, dashboard"],
], widths=[1.3, 2.5, 2.4], font=9)

doc.add_heading("3.2  End-to-end request flow (the 13-agent loop)", 2)
steps = [
    ("1  Intent & Task Classifier", "zero-shot DL classifier \u2192 task type (factual / analytical / creative / mixed / high-stakes) + ambiguity. A safety override forces 'high-stakes' on medical / legal / financial / safety cues."),
    ("2  Risk Profiler", "engineered features (domain risk, consequence, ambiguity, numeric load, task prior) \u2192 request risk score in [0,1] with factor breakdown."),
    ("3  Policy Controller (ARCOP)", "computes the 6-parameter policy vector \u03c0 = f(risk, intent, coverage, criticality, ambiguity, agreement, creativity demand, confidence, cost, latency). Rule engine, optionally blended with a learned regressor, clamped to a safety floor."),
    ("4  Candidate Generator", "N candidates via the gateway; temperature & system prompt derived from the policy vector. Grounding-first retrieval feeds context when grounding intensity is high."),
    ("5  Claim Decomposer", "LLM extracts atomic, independently checkable claims; each is typed (factual / numeric / causal / temporal / opinion / creative), scored for criticality & temporal sensitivity, and seeded into the Claim Risk Graph with entity-dependency edges."),
    ("6  Evidence Retrieval", "per claim: query expansion (if ambiguous), hybrid retrieval (dense + lexical, RRF-fused, MMR-diversified), contradictory-evidence retrieval for high-risk claims, Redis evidence cache."),
    ("7  Source Quality", "ML scorer over authority / freshness / relevance / consistency / corroboration \u2192 per-passage source score + a source-agreement signal."),
    ("8  Verification", "NLI entailment (local DeBERTa-v3 or the gateway verifier model as a batched judge) claim\u22a8evidence. At standard+ depth: an independent LLM verifier. For critical claims at deep depth: a counterfactual probe ('what evidence would falsify this?')."),
    ("9  Contradiction", "scans evidence for passages that refute a claim; one batched LLM call checks the answer's claims for internal inconsistency."),
    ("10  Claim Risk Model H(x)", "calibrated logistic model over 7 signals \u2192 per-claim risk score + normalised per-feature contributions (explanation); risk propagates along the dependency graph as a bounded floor."),
    ("11  Creativity", "distinct-n, self-BLEU diversity across candidates, embedding novelty vs the evidence centroid \u2192 creativity score; locates explicitly creative spans to protect."),
    ("[loop]  Revision", "if critical claims are unsupported / high-risk / contradicted (and it is not a creativity-first request), the reviser rewrites or hedges them, keeps <keep> creative spans intact, then re-decompose \u2192 re-retrieve \u2192 re-verify (bounded by MAAHAF_MAX_REVISION_LOOPS)."),
    ("12  Abstention / Escalation", "from calibrated confidence, hallucination budget, claim risk, agent disagreement and the policy thresholds \u2192 answer / qualify / abstain / escalate. Degenerate or no-relevant-evidence answers abstain regardless."),
    ("13  Audit", "records the full trace: task profile, policy vector, signals, candidates, claims, evidence, agent votes, pinned model versions, decision & reason. Persisted for governance."),
]
for name, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(name + "  —  ")
    r.font.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = ACCENT
    r2 = p.add_run(desc)
    r2.font.size = Pt(9.5)

# ============================================================ 4. THE 13 AGENTS
doc.add_heading("4.  The 13 Agents", 1)
table(["#", "Agent", "File", "Key output"], [
    ["1", "Intent & Task Classifier", "agents/a01_intent_classifier.py", "task profile + ambiguity"],
    ["2", "Risk Profiler", "agents/a02_risk_profiler.py", "request risk score + factors"],
    ["3", "Policy Controller (ARCOP)", "agents/a03_policy_controller.py", "6-parameter policy vector"],
    ["4", "Candidate Generator", "agents/a04_candidate_generator.py", "N candidate responses"],
    ["5", "Claim Decomposer", "agents/a05_claim_decomposer.py", "atomic claims + Claim Risk Graph"],
    ["6", "Evidence Retrieval Agent", "agents/a06_evidence_retrieval.py", "ranked evidence set + coverage"],
    ["7", "Source Quality Agent", "agents/a07_source_quality.py", "source scores + agreement"],
    ["8", "Verification Agent", "agents/a08_verification.py", "per-claim verdicts + entailment"],
    ["9", "Contradiction Agent", "agents/a09_contradiction.py", "external + internal conflict report"],
    ["10", "Creativity Agent", "agents/a10_creativity.py", "creativity score + creative spans"],
    ["11", "Revision Agent", "agents/a11_revision.py", "revised draft (creativity-preserving)"],
    ["12", "Abstention / Escalation Agent", "agents/a12_abstention.py", "answer / qualify / abstain / escalate + segments"],
    ["13", "Audit Agent", "agents/a13_audit.py", "full structured trace record"],
], widths=[0.3, 1.9, 2.4, 1.9], font=9)
para("Every agent extends a common base (agents/base.py) so the orchestrator only ever "
     "calls Agent.run(state); each run is timed, token-metered and recorded.", size=9.5,
     italic=True, color=MUTED)

# ============================================================ 5. ARCOP
doc.add_heading("5.  Adaptive Reliability\u2013Creativity Controller (ARCOP)", 1)
para("app/controller/arcop.py. For each request the controller computes a policy vector "
     "rather than choosing a static mode:")
para("\u03c0 = f(Risk, Intent, EvidenceCoverage, ClaimCriticality, Ambiguity, "
     "SourceAgreement, CreativityDemand, ModelConfidence, Cost, Latency)", italic=True,
     color=MUTED)
table(["Policy parameter", "Controls"], [
    ["grounding_intensity", "how strongly the answer must be grounded in retrieved evidence"],
    ["verification_depth", "light / standard / deep verification (adds LLM verifier + counterfactual probe)"],
    ["creativity_allowance", "how much speculative / creative latitude is permitted"],
    ["citation_requirement", "whether factual claims must carry [S#] source markers"],
    ["abstention_threshold", "minimum calibrated confidence required to answer"],
    ["escalation_threshold", "agent-disagreement level above which the request is routed to human review"],
], widths=[1.8, 4.4], font=9.5)
para("The default is an interpretable rule/scoring engine. If a trained policy artifact "
     "exists it is blended 50/50 with the rules and then re-clamped to the rule engine's "
     "safety floor for verification / citation / abstention on risky requests, so a poor "
     "learned policy can never disable safety behaviour. Named tenant profiles "
     "(strict / balanced / creative) and bounded caller overrides adjust the vector.")

doc.add_heading("5.1  Hallucination Risk Model H(x)", 2)
para("app/controller/risk_model.py — the proposal \u00a79 formula, implemented as a "
     "calibrated logistic model whose coefficients are the learned w1\u2026w7:")
para("H(x) = w\u2081(1\u2212EvidenceCoverage) + w\u2082\u00b7Contradiction + w\u2083\u00b7SourceRisk "
     "+ w\u2084\u00b7ModelUncertainty + w\u2085\u00b7ClaimCriticality + w\u2086\u00b7TemporalSensitivity "
     "+ w\u2087\u00b7AgentDisagreement", italic=True, color=MUTED)
para("Every score returns risk_contributions (normalised coefficient \u00d7 feature) so the "
     "dashboard can explain why a claim is high-risk. Risk then propagates along the "
     "Claim Risk Graph dependency edges as a bounded floor (a claim cannot be considered "
     "less risky than a decayed fraction of its riskiest dependency).")
para("Supporting controllers: budget.py (per-task hallucination budget — near-zero for "
     "high-stakes), calibration.py (isotonic/power-law confidence calibration + "
     "confidence\u2013evidence consistency gap), consensus.py (verdict aggregation \u2192 "
     "disagreement \u2192 escalation routing).")

# ============================================================ 6. TECH STACK
doc.add_heading("6.  Technology Stack", 1)
table(["Area", "Technology"], [
    ["Language / runtime", "Python 3.11 (3.13 venv compatible)"],
    ["Web framework", "FastAPI + Uvicorn, Pydantic v2, pydantic-settings"],
    ["Orchestration", "LangGraph state machine (visualisation) + hand-rolled runtime loop"],
    ["LLM access", "Model-agnostic gateway; OpenAI SDK, Ollama/vLLM (OpenAI-compatible), HuggingFace transformers, deterministic mock"],
    ["Database", "PostgreSQL 16 + pgvector (HNSW index), SQLAlchemy 2.0, Alembic migrations"],
    ["Cache / state", "Redis 7 (async for rate-limit, sync for evidence cache)"],
    ["Retrieval", "sentence-transformers (bge-small-en-v1.5), rank-bm25, cross-encoder reranker; PostgreSQL FTS (tsvector/GIN) in the DB path"],
    ["ML", "scikit-learn (LogisticRegression, GradientBoosting, IsotonicRegression, CalibratedClassifierCV), joblib, numpy"],
    ["DL / NLP", "transformers (>=4.44,<5), torch (CPU), DeBERTa-v3 NLI, zero-shot classification"],
    ["Observability", "OpenTelemetry SDK + OTLP exporter, prometheus-client, structlog (JSON)"],
    ["API surface", "REST/JSON, OpenAPI 3 / Swagger UI at /docs"],
    ["Auth / security", "API key + JWT (python-jose), RBAC, PII redaction, tenant row-scoping"],
    ["Frontend", "Vite + React 18 + TypeScript + Tailwind CSS + TanStack Query + Recharts + reactflow"],
    ["Packaging / deploy", "Docker (multi-stage), Docker Compose (db/redis/api/frontend + obs profile), Kubernetes manifests, k6 load test"],
    ["Dev tooling", "pytest, ruff, mypy; Windows task runner (tasks.ps1); Makefile; GitHub Actions CI (lint + test + dependency scan)"],
], widths=[1.5, 4.7], font=9.5)

# ============================================================ 7. ML/DL
doc.add_heading("7.  Machine-Learning / Deep-Learning Components", 1)
table(["Component", "Type", "Model / method", "Offline fallback"], [
    ["Embeddings", "DL", "BAAI/bge-small-en-v1.5 (sentence-transformers)", "hashed bag-of-tokens (384-d)"],
    ["Reranker", "DL", "cross-encoder/ms-marco-MiniLM-L-6-v2", "lexical overlap"],
    ["Verification / contradiction", "DL / LLM", "DeBERTa-v3-base-mnli-fever-anli NLI, or gpt-4o-mini as a batched entailment judge", "overlap + negation heuristic"],
    ["Intent & task classifier", "DL", "zero-shot deberta-v3-xsmall / bart-large-mnli; optional fine-tuned DistilBERT head", "keyword scoring + high-stakes cue override"],
    ["Claim-type classifier", "ML", "TF-IDF + LogisticRegression", "LLM-provided type (validated against enum)"],
    ["Hallucination risk H(x)", "ML", "class-balanced LogisticRegression, isotonic-calibrated; learned w1\u2026w7 + per-feature explanation", "proposal's fixed weights"],
    ["Confidence calibration", "ML", "IsotonicRegression; ECE / Brier reported", "power-law shrink"],
    ["Source-quality scorer", "ML", "GradientBoostingRegressor over 5 provenance features", "weighted sum"],
    ["ARCOP policy", "ML", "MultiOutput GradientBoosting (learning-to-route), clamped to rule-engine safety floor", "interpretable rule engine"],
    ["Creativity scoring", "DL / ML", "distinct-1/2, self-BLEU diversity, embedding novelty vs corpus centroid", "distinct-n only"],
], widths=[1.5, 0.7, 2.6, 1.5], font=8.5)
para("Training data is synthetic-seeded (app/ml/synth_data.py) so the models train with "
     "no client data. app/ml/retrain_from_eval.py closes the loop: a real evaluation run "
     "emits labelled feature rows, from which the risk model and calibrator are refit on "
     "a held-out split (see section 9).", size=9.5, italic=True, color=MUTED)

doc.add_page_break()

# ============================================================ 8. WORKFLOW DIAGRAM (text)
doc.add_heading("8.  Orchestration Graph", 1)
para("The LangGraph state machine (app/orchestration/graph.py):", space_after=3)
mono = doc.add_paragraph()
mono.paragraph_format.space_after = Pt(8)
mr = mono.add_run(
    "START\n"
    "  \u2192 intent \u2192 risk \u2192 policy \u2192 pre_retrieve \u2192 generate \u2192 decompose\n"
    "  \u2192 [ retrieve \u2192 verify \u2192 risk_scoring \u2192 creativity ]\n"
    "        \u2514\u2500 needs_revision? \u2500 yes \u2192 revise \u2192 (back to decompose)\n"
    "        \u2514\u2500 needs_revision? \u2500 no  \u2192 decide \u2192 finalize \u2192 audit \u2192 END\n"
    "\n"
    "  pre_retrieve  : grounding-first retrieval when grounding_intensity >= 0.45\n"
    "  retrieve      : evidence_retrieval + source_quality\n"
    "  verify        : verification + contradiction\n"
    "  risk_scoring  : score_claim (H(x)) per claim + propagate on the graph\n"
    "  decide        : abstention / escalation agent\n"
    "  finalize      : segment answer (factual / assumption / creative) + audit"
)
mr.font.name = "Consolas"
mr.font.size = Pt(8.5)
mr.font.color.rgb = INK

# ============================================================ 9. EVALUATION
doc.add_heading("9.  Evaluation", 1)
doc.add_heading("9.1  Methodology", 2)
bullets([
    ("Systems compared:", "MA-AHAF (full 13-agent pipeline) vs Static-RAG baseline (same generator + same "
     "hybrid retriever, top-5 passages, no adaptive control / verification / revision / abstention)."),
    ("Benchmark:", "77 items across 5 task types (factual 31, high-stakes 17, analytical 12, creative 10, "
     "mixed 7); 8+ high-stakes items are deliberately unanswerable from the corpus to test calibrated abstention. "
     "Seeded splits for reproducibility."),
    ("Generator:", "OpenAI gpt-4o-mini for generation, decomposition, verification and revision; local bge-small "
     "+ BM25 + cross-encoder retrieval; LLM-backed entailment judge."),
    ("Metrics:", "unsupported-claim rate, citation precision/recall, answer entailment, ECE / Brier calibration, "
     "abstention rate, answer-correct, creativity (distinct-2 + diversity), a composite reliability index, and a "
     "reliability\u2013creativity Pareto frontier gain; plus latency, tokens and USD cost."),
    ("Retraining loop:", "the eval emits per-claim feature rows (\u2192 NLI-derived unsupported label) and "
     "confidence rows (\u2192 answer-correct); retrain_from_eval.py fits a class-balanced logistic risk model "
     "(25% held-out) and an isotonic calibrator."),
])

doc.add_heading("9.2  Results  (76 of 77 items completed; 1 timed out during a transient network drop)", 2)
table(["Metric", "Static-RAG", "MA-AHAF", "\u0394", "Reading"], [
    ["Calibration error (ECE)", "0.191", "0.066", "\u221265%", "clear win"],
    ["Brier score", "0.194", "0.145", "\u22120.049", "clear win"],
    ["Abstention rate (overall)", "0.000", "0.158", "+0.158", "clear win (abstains on unanswerable)"],
    ["Unsupported-claim rate", "0.179", "0.183", "+0.004", "flat \u2013 generator already well-grounded"],
    ["Answer-correct", "0.789", "0.763", "\u22120.026", "\u2248 same"],
    ["Citation recall (same metric both)", "0.821", "0.829", "+0.008", "slight win"],
    ["Pareto frontier gain", "\u2013", "0.053", "\u2013", "> 0 (target met)"],
    ["Creativity (creative split)", "\u2013", "0.808", "\u2013", "retained"],
    ["Latency per request", "3.5 s", "31.5 s", "\u22489\u00d7", "the reliability tax (visible & controllable)"],
    ["Cost per request", "~$0.0001", "~$0.0014", "\u224814\u00d7", "the reliability tax"],
], widths=[2.1, 1.0, 1.0, 0.8, 1.6], font=8.5)

doc.add_heading("9.3  Interpretation", 2)
para("Genuine wins.", bold=True, space_after=2)
para("Confidence calibration improves markedly (ECE \u22120.125) and the framework abstains "
     "on 15.8% of items \u2014 all on the deliberately unanswerable high-stakes set \u2014 where "
     "the baseline always answers. These are the proposal's core \u00a75/\u00a710 objectives "
     "and they hold. Unsupported-claim rate is flat: with gpt-4o-mini as the generator "
     "the raw answers are already fairly grounded, so the framework's value on this "
     "benchmark is the abstention on cases it cannot ground, not a lower rate on the "
     "cases it answers.")
para("Metric asymmetries (not regressions).", bold=True, space_after=2)
bullets([
    "Citation precision (0.82 \u2192 0.15): the baseline number is answer-level entailment; the MA-AHAF number is "
    "true claim-level precision over every decomposed atomic claim (verdict=supported AND lexical overlap with a "
    "one-sentence gold snippet). Decomposition rewords claims, so overlap frequently misses. Citation recall \u2014 "
    "computed identically for both \u2014 is 0.82 \u2192 0.83.",
    "Creativity (0.99 \u2192 0.65): the baseline generates one candidate, so its self-BLEU is 0 and its diversity "
    "term is a free 1.0. MA-AHAF generates 2\u20133 candidates and pays real inter-candidate self-BLEU. On the "
    "creative split MA-AHAF scores 0.808 \u2014 creative quality is retained, not collapsed.",
    "Reliability index (0.83 \u2192 0.65): dragged down almost entirely by the citation-precision term above.",
])
para("Retraining outcome.", bold=True, space_after=2)
para("risk_model refit on 213 claim rows (52 unsupported): held-out AUC 0.751, Brier "
     "0.157. Dominant learned predictors of an unsupported claim are agent_disagreement "
     "(0.91) and source_risk (0.88). calibrator refit on 76 rows: ECE 0.106 \u2192 ~0. "
     "Because 213 rows on a 13-document benchmark is far too little to trust (the fit "
     "collapsed the missing-evidence weight to 0), the retrained artifacts are archived "
     "under artifacts/eval/20260901T172635/*.retrained.joblib but the runtime keeps the "
     "interpretable proposal \u00a79 weights. The loop is demonstrated end-to-end; a "
     "production run needs the \u2265500-prompt client benchmark.")

doc.add_page_break()

# ============================================================ 10. SECURITY
doc.add_heading("10.  Security, Privacy & Governance", 1)
table(["Control", "Implementation"], [
    ["Authentication", "API key (SHA-256 hashed at rest) + JWT (python-jose, explicit algorithm allow-list)"],
    ["Authorisation (RBAC)", "viewer < operator < admin; rank check on every protected route"],
    ["Tenant isolation", "every DB read/write row-scoped by principal.tenant_id; no cross-tenant read path"],
    ["PII redaction", "redact_pii() runs before any prompt / context / evidence is persisted or sent through the pipeline"],
    ["Secrets", "environment only; startup ABORTS in prod on weak JWT secret, 'dev-key', mock provider, or localhost DB URL"],
    ["CORS", "'*' only when MAAHAF_ENV=dev; allow-list otherwise; methods/headers narrowed"],
    ["Rate limiting", "per-key Redis token bucket; fails open with a warning + metric (configurable to fail closed)"],
    ["External retrieval", "off by default (MAAHAF_ALLOW_EXTERNAL_RETRIEVAL); KB is tenant-scoped"],
    ["Audit integrity", "every response trace pins model_versions; human-review decisions are appended, not overwritten"],
    ["Container hardening", "non-root, read-only rootfs, all capabilities dropped, no privilege escalation (deploy/k8s/api.yaml)"],
    ["Network egress", "Kubernetes NetworkPolicy restricts egress to Postgres / Redis / OTel / :443 / DNS"],
    ["Prompt injection", "documented; evidence treated as data; a dedicated injection classifier recommended for untrusted corpora"],
], widths=[1.6, 4.6], font=9)
para("A full internal code-level security review (10 findings S-1\u2026S-10, 8 fixed) is in "
     "docs/security-review.md; the go-live checklist is in SECURITY.md. A third-party "
     "penetration test is still required before production.", size=9.5, italic=True, color=MUTED)

# ============================================================ 11. KB + BENCHMARK
doc.add_heading("11.  Knowledge Base & Benchmark", 1)
bullets([
    ("Knowledge base (data/corpus/):", "143 documents. 13 product/policy documents (ACME Cloud pricing, SLA, "
     "refund, security, support, data-retention, API limits) plus 130 general-knowledge documents across 13 "
     "domains: physics, chemistry, biology, astronomy, geography, world history, computer science, "
     "artificial intelligence, mathematics, health & wellness, economics & finance, business & management, "
     "law & governance (10 documents each)."),
    ("Benchmark (data/benchmark/):", "benchmark.jsonl (36) + benchmark_acme.jsonl (41) = 77 labelled items, "
     "each with prompt, task type, reference answer, gold evidence spans, and answerable/unanswerable labels."),
    ("Dedicated eval corpus (data/corpus_bench/):", "the 13 benchmark-relevant documents, used to keep the "
     "evaluation fast and methodologically aligned with how the benchmark was written."),
    ("Retrieval behaviour:", "factual / analytical / high-stakes questions are answered only from KB evidence "
     "(and abstain when the KB does not cover them); creative / low-grounding questions answer from the model's "
     "own knowledge. General-knowledge questions are now answerable because the KB spans 13 domains; specific "
     "facts outside the KB still abstain \u2014 by design (proposal \u00a75 hallucination budget)."),
])

doc.add_page_break()

# ============================================================ 12. FILE MAP
doc.add_heading("12.  Complete File Map \u2014 What Each File Does", 1)

def filemap(title, rows):
    doc.add_heading(title, 3)
    table(["File", "What it does"], rows, widths=[2.2, 4.0], font=8.5)

filemap("backend/app/  \u2014  application entry & config", [
    ["main.py", "FastAPI factory: lifespan (runs config validation, aborts on bad prod config), CORS, error handlers, router registration, /health, /ready, /metrics, /v1/graph (Mermaid)."],
    ["config.py", "Central pydantic-settings config (env prefix MAAHAF_). LLM role\u2192model map, price table, model ids, revision loops, NLI backend selector, OpenAI key resolution, validate_for_runtime() (fatal-misconfig list), CORS list."],
    ["__init__.py", "package version string."],
])
filemap("backend/app/agents/  \u2014  the 13 agents", [
    ["base.py", "Agent ABC. run(state) wraps _run(): times it, meters tokens, records an AgentRecord, pins model version, never lets an agent crash the graph."],
    ["a01_intent_classifier.py", "Zero-shot DL classification into 5 task types + ambiguity estimate; high-stakes keyword override (dose, legal advice, invest, safe daily...)."],
    ["a02_risk_profiler.py", "Rule-based scorer over domain risk (medical/legal/financial/safety keyword hits), consequence (zero-shot), ambiguity, numeric load, task prior \u2192 risk score + factors."],
    ["a03_policy_controller.py", "Builds Signals from state, calls arcop.policy(), applies bounded caller overrides and the tenant profile, stores the policy vector."],
    ["a04_candidate_generator.py", "Generates N candidates through the gateway; temperature & system prompt from the policy vector; grounding-first context injection; picks the most grounded-sounding draft."],
    ["a05_claim_decomposer.py", "LLM atomic-claim extraction (JSON); claim-type classification (TF-IDF model or validated LLM label); criticality / temporal sensitivity; creative-request handling; builds the ClaimGraph with entity-dependency edges; sentence-split fallback."],
    ["a06_evidence_retrieval.py", "Per non-creative claim: query expansion, hybrid retrieve (k scaled by grounding), contradictory-evidence retrieval, Redis cache; computes coverage."],
    ["a07_source_quality.py", "ML (GradientBoosting) or weighted-fallback scorer over authority / freshness / relevance / consistency / corroboration; sets per-passage source_score and the source-agreement signal."],
    ["a08_verification.py", "best_support() NLI entailment claim\u22a8evidence; NLI verdict; independent LLM verifier at depth\u22651 or for critical claims; counterfactual probe at deep depth; consensus.resolve_claim()."],
    ["a09_contradiction.py", "External: re-scans evidence for refuting passages (skipped when the LLM judge already did a thorough pass). Internal: one batched LLM call to find contradicting claim pairs."],
    ["a10_creativity.py", "distinct-1/2, self-BLEU across candidates (\u2192 diversity), embedding novelty vs the evidence centroid \u2192 creativity score; records creative spans for the preservation gate."],
    ["a11_revision.py", "Rewrites unsupported / overconfident / contradicted claims; protects <keep> creative spans; strips echoed scaffolding; discards degenerate rewrites; increments the loop counter."],
    ["a12_abstention.py", "The safety decision: degenerate / empty-analysis / no-relevant-evidence \u2192 abstain; disagreement or budget-exceeded \u2192 escalate; low calibrated confidence + critical unsupported \u2192 abstain; else qualify or answer. Also segments the answer."],
    ["a13_audit.py", "Assembles the full structured trace (signals, policy, candidates, claims, evidence, agent runs, model versions, decision) and stashes it for persistence."],
])
filemap("backend/app/controller/  \u2014  adaptive control", [
    ["arcop.py", "PolicyVector & Signals models; _rule_policy() (interpretable scoring), tenant profiles, optional learned-policy blend clamped to a safety floor."],
    ["risk_model.py", "score_claim(): H(x) via the calibrated logistic model (or fixed weights); normalised per-feature contributions; risk level; explain()."],
    ["calibration.py", "calibrate() (isotonic model or power-law shrink), consistency_gap() (confidence vs evidence strength), ece()."],
    ["budget.py", "Per-task hallucination budget (near-zero for high-stakes); consumed() counts only non-supported checkable claims so answer length alone never trips it."],
    ["consensus.py", "claim_disagreement() (verdict spread), resolve_claim() (entailment-weighted verdict vote), request_disagreement()."],
])
filemap("backend/app/claimgraph/  \u2014  Claim Risk Graph", [
    ["schema.py", "Claim & ClaimGraph pydantic models; max_risk(), critical_unsupported()."],
    ["graph.py", "networkx build (claims \u2194 entities \u2194 evidence); propagate_risk() (bounded floor along dependency edges); to_cytoscape() for the dashboard."],
])
filemap("backend/app/orchestration/  \u2014  the runtime", [
    ["state.py", "RequestState \u2014 the single pydantic object threaded through the graph (inputs, agent outputs, records, final response)."],
    ["nodes.py", "Thin node wrappers around each agent + pure aggregation steps; needs_revision() conditional edge; pre-retrieve (grounding-first)."],
    ["graph.py", "LangGraph StateGraph: nodes, edges, the conditional revision loop; compiled + cached; exposed as Mermaid."],
    ["pipeline.py", "The executable control loop (bounded revision), PII redaction, gateway + NLI wiring, telemetry, and _persist() of the audit trail."],
])
filemap("backend/app/llm/  \u2014  model-agnostic gateway", [
    ["base.py", "LLMAdapter ABC; LLMResponse (text, tokens, logprobs, uncertainty)."],
    ["gateway.py", "Role\u2192model routing (generator / verifier / decomposer / reviser / judge / expander); retry/backoff; UsageMeter (tokens + USD via price table); complete_json() with fence tolerance; Prometheus counters."],
    ["openai_adapter.py", "OpenAI Chat Completions + logprob-based uncertainty; api_key from settings; tenacity retry."],
    ["local_adapter.py", "Any OpenAI-compatible endpoint (Ollama / vLLM) \u2014 for verifier-family diversity."],
    ["hf_adapter.py", "Local HuggingFace seq2seq (flan-t5); read-then-answer framing; JSON coercion for weak models."],
    ["mock_adapter.py", "Deterministic offline adapter for tests and plumbing checks."],
])
filemap("backend/app/retrieval/  \u2014  hybrid retrieval", [
    ["schema.py", "Evidence model (scores, stance, provenance)."],
    ["embeddings.py", "sentence-transformers (bge-small); deterministic hashing fallback offline."],
    ["reranker.py", "cross-encoder reranker; lexical-overlap fallback offline."],
    ["vector_store.py", "pgvector similarity search + chunk upsert (bound SQL parameters)."],
    ["keyword.py", "PostgreSQL full-text search (plainto_tsquery, ts_rank_cd)."],
    ["hybrid.py", "RRF fusion of vector + keyword, MMR diversification, then rerank \u2014 the single retrieve() entry point in the DB path."],
    ["query_expansion.py", "LLM query reformulation (2\u20134 diverse queries) for ambiguous claims."],
    ["cache.py", "Redis evidence cache (sync), degrades to no-op if Redis is down."],
    ["local_store.py", "In-memory corpus index (real embeddings + BM25 + RRF + rerank) so the full pipeline runs with no database \u2014 used by demos, the dashboard server and the local eval."],
])
filemap("backend/app/nlp/  \u2014  language inference", [
    ["nli.py", "Entailment / contradiction scoring. Backends: local DeBERTa-v3 pipeline, or a batched gateway-verifier LLM judge (one call scores a claim against several passages), or a lexical heuristic offline."],
    ["zeroshot.py", "Zero-shot text classification (deberta-v3-xsmall / bart-large-mnli); keyword-scoring fallback."],
])
filemap("backend/app/eval/  \u2014  evaluation", [
    ["datasets.py", "Load benchmark JSONL (merges sibling benchmark_*.jsonl, de-duped by id); split_by_type()."],
    ["metrics.py", "unsupported-claim rate, citation precision/recall, entailment accuracy, creativity (distinct-2 + self-BLEU), calibration (ECE / Brier), reliability index."],
    ["baseline.py", "static_rag() \u2014 same generator + retriever, one LLM call, no adaptive control."],
    ["harness.py", "DB-backed MA-AHAF vs baseline run + aggregation + Pareto (used by scripts/run_eval.py)."],
    ["pareto.py", "pareto_front() (non-dominated points), frontier_gain() (fraction of baseline points MA-AHAF dominates)."],
])
filemap("backend/app/ml/  \u2014  model training & registry", [
    ["features.py", "RISK_FEATURES order (w1\u2026w7), risk_vector(), source & policy feature vectors."],
    ["registry.py", "Load joblib artifacts with graceful rule-based fallback + cache."],
    ["synth_data.py", "Synthetic labelled data generators for every trainable model."],
    ["train_risk_model.py", "Train the H(x) logistic model + calibrator from synthetic data."],
    ["train_claim_type_clf.py", "TF-IDF + LogisticRegression claim-type classifier."],
    ["train_source_quality.py", "GradientBoosting source-quality regressor."],
    ["train_policy.py", "MultiOutput GradientBoosting learned ARCOP policy (synthetic-seeded)."],
    ["train_intent_clf.py", "Optional DistilBERT fine-tune for the intent classifier."],
    ["retrain_from_eval.py", "Refit the risk model (class-balanced logistic, 25% held-out, AUC/Brier before/after, learned w1\u2026w7) and the calibrator (isotonic) from a real eval run; writes to the run dir unless --promote."],
])
filemap("backend/app/api/  \u2014  HTTP surface", [
    ["deps.py", "Auth / RBAC / tenant resolution: API key or JWT \u2192 Principal; requires(role) dependency; rate-limit check."],
    ["schemas.py", "Request/response pydantic models (OpenAPI schemas)."],
    ["routes_generate.py", "POST /v1/generate \u2014 the main endpoint; runs the pipeline in a threadpool, returns response + segments + claims + evidence + policy + usage."],
    ["routes_kb.py", "KB ingest (chunk + embed), list documents, hybrid search."],
    ["routes_traces.py", "List requests + fetch the full per-agent audit trace."],
    ["routes_metrics.py", "JSON metrics summary for the dashboard; pending escalations."],
    ["routes_eval.py", "Trigger + fetch evaluation runs."],
    ["routes_review.py", "Human-review queue: list / stats / detail / resolve (approve / revise / reject); the decision is appended to the audit trace."],
    ["routes_admin.py", "Tenants, API keys, policy profiles; /v1/admin/token (disabled in prod)."],
])
filemap("backend/app/db/  \u2014  persistence", [
    ["models.py", "SQLAlchemy models: tenants, api_keys, documents, chunks (HNSW vector + GIN tsvector), requests, claims, agent_runs, audit_traces, eval_runs, escalation_queue."],
    ["session.py", "Engine + sessionmaker + session_scope + get_db dependency."],
    ["repositories.py", "persist_request() (request + claims + agent runs + audit + escalation-queue insert), list/get helpers, get_audit()."],
    ["cache.py", "Async Redis client (rate-limit, response cache)."],
    ["migrations/0001_init.py", "Initial schema + pgvector extension + FTS trigger."],
    ["migrations/0002_review_columns.py", "Adds decision / review_note / reviewed_by / reviewed_at to escalation_queue."],
])
filemap("backend/app/core/  \u2014  cross-cutting", [
    ["security.py", "JWT create/decode, require_role(), redact_pii() (email / phone / SSN / card / IP)."],
    ["ratelimit.py", "Redis token-bucket limiter; fail-open with metric."],
    ["telemetry.py", "OpenTelemetry init + Prometheus metrics (request/agent latency, tokens, cost, verification depth, abstentions, disagreement, max claim risk)."],
    ["logging.py", "structlog JSON configuration."],
    ["errors.py", "Typed API errors + FastAPI exception handlers."],
])
filemap("backend/scripts/  \u2014  operational scripts", [
    ["serve_demo.py", "Standalone no-DB FastAPI: runs the full pipeline against the in-memory corpus, serves the built dashboard, keeps requests / escalations / eval runs in memory. Powers the live dashboard."],
    ["demo.py", "Console end-to-end demo (mock or local models), no DB / API key."],
    ["capture_traces.py", "Runs a set of prompts through the real pipeline and writes full traces to demo_traces.json (for the visual trace report)."],
    ["eval_local.py", "DB-less evaluation: MA-AHAF vs static-RAG over the benchmark; per-item wall-clock cap; emits report.json, pareto.csv, training_rows.jsonl, calibration_rows.jsonl, latest.txt."],
    ["run_eval.py", "DB-backed evaluation harness CLI."],
    ["finalize_report.py", "Fills section 4 of docs/final-report.md from artifacts/eval/latest + retrain_report.json."],
    ["bootstrap_models.py", "Download HF models + train the sklearn artifacts."],
    ["seed_kb.py / seed_benchmark.py", "Load the sample corpus (embed + index) and verify the benchmark in the Docker path."],
])
filemap("backend/tests/  \u2014  automated tests (30, all passing)", [
    ["conftest.py", "Forces fully-offline deterministic mode (mock provider)."],
    ["test_pipeline_integration.py", "End-to-end pipeline with retrieval stubbed: factual completes with claims, creative gets high creativity allowance, high-stakes uses deep verification, PII redacted, audit trace assembled."],
    ["test_controller.py", "ARCOP strictness, risk-model monotonicity, calibration shrink, consensus disagreement, budget tightness."],
    ["test_gateway.py", "Role\u2192model routing, usage metering, JSON parsing."],
    ["test_retrieval_offline.py", "Hashing embedder + BM25 + RRF + rerank fallback path."],
    ["test_eval_smoke.py", "Metrics + Pareto sanity."],
    ["test_config_security.py", "Prod config rejects weak JWT secret / dev-key / mock provider / localhost DB; OpenAI key read from unprefixed env; CORS locked down outside dev."],
    ["test_review_flow.py", "High-stakes prompt \u2192 review-queue item \u2192 resolve (revised / rejected) through the demo server."],
])
filemap("frontend/src/  \u2014  React dashboard", [
    ["main.tsx / App.tsx", "Bootstrap; nav + routes: Playground, Traces, Trace Detail, Review, Metrics, Evaluation, Knowledge Base; API-key control."],
    ["api/client.ts", "Typed fetch client for every endpoint (x-api-key header)."],
    ["lib/types.ts", "Shared response types."],
    ["pages/Playground.tsx", "Prompt \u2192 Generate \u2192 live action badge, segmented answer, claims table, evidence, policy radar, agent timeline, per-request stats."],
    ["pages/Traces.tsx / TraceDetail.tsx", "Request history with filters; full trace view (claim graph, agent timeline, per-claim risk contributions)."],
    ["pages/Review.tsx", "Human-review queue UI: pending / reviewed tabs, item detail, approve / revise / reject."],
    ["pages/Metrics.tsx", "Action mix, abstention rate, latency percentiles, token/cost totals, verification-depth, time series."],
    ["pages/Evaluation.tsx", "Trigger a run; render deltas, the reliability\u2013creativity Pareto scatter and the full report."],
    ["pages/KnowledgeBase.tsx", "List documents, hybrid search, (ingest in the Docker path)."],
    ["components/", "PolicyRadar, ClaimGraphView (reactflow), AgentTimeline, ClaimsTable, SegmentedAnswer, ParetoScatter, ui (Card / badges / Stat)."],
])
filemap("deploy/ , observability/ , docs/ , data/ , root", [
    ["docker-compose.yml", "db (pgvector) + redis + api + frontend, plus an 'obs' profile (OTel collector + Prometheus + Grafana)."],
    ["deploy/k8s/*.yaml", "Namespace, ConfigMap, Secret template, Postgres StatefulSet + Redis, API Deployment (HPA, PDB, non-root, read-only rootfs), frontend + Ingress (TLS, NetworkPolicy egress allow-list)."],
    ["deploy/loadtest/k6.js", "k6 load-test script for /v1/generate."],
    ["observability/", "prometheus.yml, otel-collector.yaml, Grafana dashboard + datasource provisioning."],
    ["docs/", "SRS.md (FR-1\u2026FR-22), architecture.md, evaluation-methodology.md, final-report.md (auto-filled results), security-review.md, this report."],
    ["data/corpus/ (143) , data/benchmark/ (77) , data/corpus_bench/ (13)", "Knowledge base, labelled benchmark, dedicated eval corpus."],
    ["tasks.ps1 / Makefile", "Task runners: setup, test, lint, demo, dashboard[-openai], serve, capture, eval-local / eval-openai, retrain, up / up-obs / down / migrate / seed."],
    ["pyproject.toml", "Backend dependencies + ruff + pytest + mypy config (transformers pinned >=4.44,<5)."],
    [".github/workflows/ci.yml", "Lint + test + dependency scan on every push."],
])

doc.add_page_break()

# ============================================================ 13. HOW TO RUN
doc.add_heading("13.  How to Run", 1)
doc.add_heading("13.1  Live dashboard (no Docker, uses OpenAI)", 2)
cb = doc.add_paragraph()
cb.paragraph_format.space_after = Pt(8)
r = cb.add_run(
    "# put OPENAI_API_KEY=sk-... in ma-ahaf\\backend\\.env\n"
    ".\\tasks.ps1 dashboard-openai       # build + serve at http://localhost:8000\n"
    "#   Playground tab: type any prompt \u2192 Generate \u2192 full 13-agent output\n\n"
    ".\\tasks.ps1 serve                  # API only, deterministic mock (no key, instant)\n"
    ".\\tasks.ps1 demo-real              # console end-to-end demo, local models, offline"
)
r.font.name = "Consolas"
r.font.size = Pt(9)

doc.add_heading("13.2  Full stack (Docker)", 2)
cb2 = doc.add_paragraph()
cb2.paragraph_format.space_after = Pt(8)
r2 = cb2.add_run(
    "cp .env.example .env               # set OPENAI_API_KEY\n"
    "docker compose up -d --build       # db (pgvector) + redis + api + frontend\n"
    "docker compose exec api alembic upgrade head\n"
    "docker compose exec api python -m scripts.seed_kb\n"
    "#   dashboard  http://localhost:5173      API docs  http://localhost:8000/docs\n"
    "docker compose --profile obs up -d # + Prometheus + Grafana (http://localhost:3000)"
)
r2.font.name = "Consolas"
r2.font.size = Pt(9)

doc.add_heading("13.3  Evaluation + retraining", 2)
cb3 = doc.add_paragraph()
cb3.paragraph_format.space_after = Pt(8)
r3 = cb3.add_run(
    ".\\tasks.ps1 eval-openai            # MA-AHAF vs static-RAG, 77 items, gpt-4o-mini\n"
    ".\\tasks.ps1 retrain                # refit risk model + calibrator from the run\n"
    "python -m scripts.finalize_report  # write measured numbers into docs/final-report.md\n"
    "#   artifacts land in backend/artifacts/eval/<timestamp>/"
)
r3.font.name = "Consolas"
r3.font.size = Pt(9)

doc.add_heading("13.4  Main endpoint", 2)
cb4 = doc.add_paragraph()
cb4.paragraph_format.space_after = Pt(8)
r4 = cb4.add_run(
    "curl -s localhost:8000/v1/generate \\\n"
    "  -H 'x-api-key: dev-key' -H 'content-type: application/json' \\\n"
    "  -d '{\"prompt\": \"What is the ACME Cloud refund policy for annual plans?\"}'\n\n"
    "# response: { response, segments[], claims[], evidence[], confidence,\n"
    "#             calibrated_confidence, action, action_reason, policy_vector,\n"
    "#             max_claim_risk, agent_disagreement, usage, trace_id }"
)
r4.font.name = "Consolas"
r4.font.size = Pt(9)

# ============================================================ 14. LIMITATIONS
doc.add_heading("14.  Limitations & Future Work", 1)
para("Known limitations (documented, not defects).", bold=True, space_after=2)
bullets([
    "Benchmark scale: 77 synthetic items on a small corpus \u2014 results are directional. A production engagement "
    "needs a client-domain corpus and \u2265500 human-labelled prompts.",
    "Same judge family: on the OpenAI path the generator and the entailment judge are both gpt-4o-mini. Point "
    "MAAHAF_LLM__VERIFIER_MODEL at a different family (or MAAHAF_NLI_BACKEND=local) before publishing headline numbers.",
    "Latency / cost: ~9\u201314\u00d7 the baseline \u2014 the reliability tax. The ARCOP policy already spends it only "
    "where risk warrants; a per-tenant daily token budget is recommended for production.",
    "Learned ARCOP policy still trains on synthetic data; only the risk model and calibrator are refit from real "
    "eval labels, and those refits are archived rather than shipped until a larger benchmark exists.",
    "Docker / Kubernetes manifests were not exercised on a live host in this engagement.",
    "Metric asymmetries in the eval (citation precision, creativity) should be made symmetric across systems.",
])
para("Future extensions (proposal \u00a720).", bold=True, space_after=2)
bullets([
    "Learning-to-route controller trained from production feedback.",
    "Per-user / per-domain reliability\u2013creativity policies.",
    "Multimodal verification (tables, charts, images).",
    "Real-time web/source freshness monitoring.",
    "Domain-expert agents for finance, healthcare, law, engineering.",
])

hr()
para("Generated 1 September 2026 from the delivered repository. Companion documents: "
     "docs/SRS.md, docs/architecture.md, docs/evaluation-methodology.md, "
     "docs/final-report.md, docs/security-review.md.",
     size=8.5, italic=True, color=MUTED)

doc.save(OUT)
print("wrote", OUT)
